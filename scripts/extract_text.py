"""
extract_text.py - 多格式文字抽取腳本
支援：.docx, .pdf, .xlsx/.xls, .jpg/.jpeg/.png
輸出 JSON：{text, metadata, ocr_used, page_count}
"""
import hashlib
import json
import os
import sys

# Force UTF-8 output on Windows to avoid cp950 mojibake
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")


def calc_hash(file_path: str) -> str:
    """計算檔案 SHA-256 hex 字串"""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_word(path: str) -> tuple[str, int]:
    """從 .docx 抽取段落與表格文字，回傳 (text, page_count)"""
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts), len(doc.sections) or 1


def _ocr_pdf_page(args):
    """Worker: render + OCR a single PDF page. Top-level so multiprocessing can pickle it.

    Each worker independently renders its own page (via poppler) — avoids large IPC
    transfer of pre-rendered images. PIL/pytesseract imports are inside the worker
    so they're loaded once per worker process.
    """
    import pytesseract
    from pdf2image import convert_from_path
    pdf_path, page_idx, dpi = args
    images = convert_from_path(pdf_path, dpi=dpi, first_page=page_idx + 1, last_page=page_idx + 1)
    if not images:
        return page_idx, ""
    text = pytesseract.image_to_string(images[0], lang="chi_tra+eng", config="--oem 1")
    return page_idx, text


def _console_python() -> str:
    """Return a console python.exe path for multiprocessing children.

    api_server.py runs under pythonw.exe (no console). multiprocessing under
    pythonw.exe DEADLOCKS on Windows — the spawned children have no valid
    stdin/stdout/stderr, so the spawn bootstrap hangs forever (this caused the
    stuck OCR process trees). Point the workers at the console python.exe instead.
    """
    exe = sys.executable or ""
    if os.path.basename(exe).lower() == "pythonw.exe":
        cand = os.path.join(os.path.dirname(exe), "python.exe")
        if os.path.isfile(cand):
            return cand
    return exe


def _ocr_pdf_sequential(path: str, n_pages: int, dpi: int) -> str:
    """OCR pages one-by-one in the current process. Slower but never deadlocks —
    used as a fallback when the parallel pool fails or times out."""
    import pytesseract
    from pdf2image import convert_from_path
    parts = []
    for i in range(n_pages):
        images = convert_from_path(path, dpi=dpi, first_page=i + 1, last_page=i + 1)
        if images:
            parts.append(pytesseract.image_to_string(images[0], lang="chi_tra+eng", config="--oem 1"))
    return "\n".join(parts)


def extract_pdf_ocr(path: str) -> tuple[str, int]:
    """OCR each PDF page in parallel (one worker process per page).

    Pages are rendered + OCR'd independently in worker processes via
    multiprocessing.Pool, leveraging multi-core CPUs. On 20-core hardware
    ≈10× speedup vs sequential, scales near-linearly with page count up to
    cpu_count. Falls back to sequential OCR if the pool errors or times out
    (e.g. multiprocessing fragility under pythonw.exe).
    """
    import multiprocessing as mp
    from pdf2image import pdfinfo_from_path

    info = pdfinfo_from_path(path)
    n_pages = info.get("Pages", 0)
    if n_pages == 0:
        return "", 0

    DPI = 200

    if n_pages == 1:
        # Single-page: skip pool spawn overhead entirely.
        return _ocr_pdf_sequential(path, 1, DPI), 1

    ctx = mp.get_context("spawn")
    try:
        ctx.set_executable(_console_python())
    except Exception:
        pass

    n_workers = min(ctx.cpu_count(), n_pages)
    args = [(path, i, DPI) for i in range(n_pages)]
    # Generous overall ceiling so a wedged worker can't hang forever; the
    # api_server subprocess timeout is the outer backstop.
    pool_timeout = max(180, n_pages * 60)

    pool = ctx.Pool(processes=n_workers)
    try:
        results = pool.map_async(_ocr_pdf_page, args).get(timeout=pool_timeout)
        pool.close()
        pool.join()
        results.sort(key=lambda x: x[0])
        return "\n".join(r[1] for r in results), n_pages
    except Exception as e:
        pool.terminate()
        pool.join()
        sys.stderr.write(f"[extract_text] parallel OCR failed ({e!r}); falling back to sequential\n")
        return _ocr_pdf_sequential(path, n_pages, DPI), n_pages


def _parse_area_table_rows(ocr_text: str) -> list[str]:
    """從 OCR 文字中偵測面積計算表格行，每戶轉成結構化文字。

    偵測條件：至少 5 行符合「戶號 + 數字序列」格式。
    欄位對應（12欄制）：
      欄1-8   建物各層面積(m²)
      欄9     總面積(m²)
      欄10    室內面積(m²)
      欄11    公設面積(m²)
      欄12    總坪數(坪)
    回傳空 list 表示非表格型 PDF，應使用一般切片。
    """
    import re

    # 修正 OCR 常見錯誤：字母O→0、字母l→1（在數字序列中）
    def fix_ocr_nums(s: str) -> str:
        s = re.sub(r'(?<=[A-Z])O(?=\d)', '0', s)   # AO66 → A066 (unit prefix)
        s = re.sub(r'(?<=\d)O(?=\d)', '.', s)       # 4O66 → 4.66
        s = re.sub(r'(?<=\d)O(?=\s)', '.', s)
        return s

    # 找出符合「戶號 數字...」的行
    # 允許：A1 A2 ... A16、Al (OCR誤讀A1)、地下室 等
    row_re = re.compile(
        r'^(A\d{1,2}|Al|地下室|小計)\s+([\d.O]+(?:\s+[\d.O]+){7,})',
        re.MULTILINE,
    )

    matches = list(row_re.finditer(fix_ocr_nums(ocr_text)))
    if len(matches) < 3:
        return []  # 非表格型 PDF

    COL_LABELS_12 = [
        "建物面積1", "建物面積2", "建物面積3", "建物面積4",
        "建物面積5", "建物面積6", "建物面積7", "建物面積8",
        "總面積m²", "室內面積m²", "公設面積m²", "總坪數",
    ]
    COL_LABELS_10 = [
        "建物面積1", "建物面積2", "建物面積3", "建物面積4",
        "建物面積5", "建物面積6",
        "總面積m²", "室內面積m²", "公設面積m²", "總坪數",
    ]

    structured = []
    for m in matches:
        unit = m.group(1).replace("Al", "A1")  # OCR 字母l修正
        vals_raw = re.split(r'\s+', m.group(2).strip())
        vals = []
        for v in vals_raw:
            try:
                vals.append(float(v))
            except ValueError:
                pass  # 略過無法解析的 OCR 殘留字元

        n = len(vals)
        if n == 12:
            labels = COL_LABELS_12
        elif n == 10:
            labels = COL_LABELS_10
        else:
            # 欄數不符：直接標 最後一欄為總坪數
            labels = [f"欄{i+1}" for i in range(n)]
            labels[-1] = "總坪數"
            if n >= 4:
                labels[-2] = "公設面積m²"
                labels[-3] = "室內面積m²"
                labels[-4] = "總面積m²"

        pairs = ", ".join(f"{labels[i]}={vals[i]}" for i in range(min(n, len(labels))))
        structured.append(f"{unit}戶 面積資料: {pairs}")

    return structured


def extract_pdf(path: str) -> tuple[str, bool, int]:
    """從 PDF 抽取文字；若為空則改用 OCR；若偵測到面積表格則結構化輸出。
    回傳 (text, ocr_used, page_count)。
    text 中若含 \\n\\n 分隔的結構化行，batch_ingest 會逐行切片。
    """
    from pypdf import PdfReader
    reader = PdfReader(path)
    page_count = len(reader.pages)
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    pypdf_text = "\n".join(parts).strip()

    # 若 pypdf 只得到表單欄位名稱（< 100 字且無中文），強制 OCR
    import re as _re
    has_cjk = bool(_re.search(r'[\u4e00-\u9fff]', pypdf_text))
    if not pypdf_text or (len(pypdf_text) < 100 and not has_cjk):
        ocr_text, ocr_pages = extract_pdf_ocr(path)
        rows = _parse_area_table_rows(ocr_text)
        if rows:
            return "\n\n".join(rows), True, max(page_count, ocr_pages)
        return ocr_text, True, max(page_count, ocr_pages)

    # 嘗試從 pypdf 文字偵測面積表格
    rows = _parse_area_table_rows(pypdf_text)
    if rows:
        return "\n\n".join(rows), False, page_count

    return pypdf_text, False, page_count


def extract_excel(path: str) -> tuple[str, int]:
    """從 .xlsx/.xls 讀取所有 sheet 文字，保留 sheet_name 前綴"""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_parts = [f"[Sheet: {sheet_name}]"]
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                sheet_parts.append(row_text)
        parts.append("\n".join(sheet_parts))
    wb.close()
    return "\n\n".join(parts), len(wb.sheetnames)


def extract_image(path: str) -> tuple[str, int]:
    """對圖片執行 OCR（需 Tesseract）"""
    import pytesseract
    from PIL import Image
    img = Image.open(path)
    text = pytesseract.image_to_string(img, lang="chi_tra+eng")
    return text, 1


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python extract_text.py <file_path>"}, ensure_ascii=False))
        sys.exit(1)

    file_path = os.path.abspath(sys.argv[1])
    if not os.path.isfile(file_path):
        print(json.dumps({"error": f"File not found: {file_path}"}, ensure_ascii=False))
        sys.exit(1)

    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    file_size = os.path.getsize(file_path)
    hash_sha256 = calc_hash(file_path)

    text = ""
    ocr_used = False
    page_count = 1

    try:
        if file_ext == ".docx":
            text, page_count = extract_word(file_path)
        elif file_ext == ".pdf":
            text, ocr_used, page_count = extract_pdf(file_path)
        elif file_ext in (".xlsx", ".xls"):
            text, page_count = extract_excel(file_path)
        elif file_ext in (".jpg", ".jpeg", ".png"):
            text, page_count = extract_image(file_path)
            ocr_used = True
        else:
            print(json.dumps({"error": f"Unsupported file type: {file_ext}"}, ensure_ascii=False))
            sys.exit(1)
    except Exception as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))
        sys.exit(1)

    result = {
        "text": text,
        "metadata": {
            "file_name": file_name,
            "file_ext": file_ext,
            "file_path": file_path,
            "file_size": file_size,
            "hash_sha256": hash_sha256,
        },
        "ocr_used": ocr_used,
        "page_count": page_count,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    import multiprocessing
    multiprocessing.freeze_support()
    main()
